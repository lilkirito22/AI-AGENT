import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from tkinter.ttk import Progressbar
import threading

from dotenv import load_dotenv

from pydantic import BaseModel
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser, StrOutputParser
from langchain.agents import create_tool_calling_agent, AgentExecutor

from docx import Document
from PyPDF2 import PdfReader
import openpyxl
import csv


# Carregar variáveis de ambiente do arquivo .env
"""
LEMBRESE de criar um arquivo .env na raiz do projeto com a variável GEMINI_API_KEY ou da IA que vc preferir  definida
tambem lembre-se de nao compartilhar esse arquivo .env com ninguém, pois ele contém informações sensíveis.
e caso va publicar no github, adicione o .env no seu .gitignore

"""
load_dotenv()


class ResearchResponse(BaseModel):
    topic: str
    summary: str
    sources: list[str]
    tools_used: list[str]


llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")

parser = PydanticOutputParser(pydantic_object=ResearchResponse)

# Cria um prompt template para categorização
prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            """Você é um assistente de organização de arquivos extremamente inteligente e autônomo.
     Sua tarefa é analisar o conteúdo ou a descrição de um arquivo e identificar a CATEGORIA MAIS ESPECÍFICA E RELEVANTE para ele.
     A categoria deve ser concisa (1-3 palavras, no máximo) e refletir o tema principal ou propósito do arquivo.
     Pense em como um humano organizaria esses arquivos em pastas lógicas.
     Exemplos de categorias esperadas: "Relatórios Financeiros", "Contratos Legais", "Receitas Culinárias", "Fotos de Viagem", "Material de Estudo", "Projetos Pessoais", "Documentos de Identidade", "Faturas de Contas", "Música Favorita", "Downloads Aleatórios", "Códigos Fonte", "Scripts", "Planilhas de Dados", "Configurações", "Arquivos de Sistema".
     Não inclua aspas, pontuação extra, ou qualquer outra explicação na sua resposta, apenas a categoria.
     Se o conteúdo for ambíguo, genérico ou insuficiente, use uma categoria ampla como "Vários" ou "Geral".""",
        ),
        ("user", "Conteúdo/Descrição do arquivo:\n---\n{file_content}\n---"),
    ]
)

# Cria uma "chain" simples: prompt -> modelo -> parser de saída
# O StrOutputParser apenas garante que a saída seja uma string limpa.
if llm:
    categorization_chain = prompt | llm | StrOutputParser()
else:
    categorization_chain = None


# Função para extrair texto de arquivos
def extrair_texto_do_arquivo(caminho_arquivo):
    _, extensao = os.path.splitext(caminho_arquivo)
    extensao = extensao.lower()
    nome_arquivo = os.path.basename(caminho_arquivo)

    # Tipos que podem ser lidos como texto puro
    text_extensions = (
        ".txt",
        ".log",
        ".md",
        ".json",
        ".xml",
        ".yaml",
        ".yml",
        ".ini",
        ".cfg",
        ".py",
        ".java",
        ".c",
        ".cpp",
        ".h",
        ".hpp",
        ".cs",
        ".js",
        ".ts",
        ".html",
        ".css",
        ".scss",
        ".sh",
        ".bat",
        ".ps1",
        ".php",
        ".go",
        ".rb",
        ".swift",
        ".kt",
        ".r",
        ".sql",
        ".asm",
        ".circ",  # Adicionado .asm e .circ
    )

    if extensao in text_extensions:
        try:
            with open(caminho_arquivo, "r", encoding="utf-8") as f:
                content = f.read()
                # Limita o tamanho para não sobrecarregar o Gemini com arquivos muito grandes
                return f"Nome do arquivo: {nome_arquivo}\nConteúdo (parcial): {content[:5000]}"
        except UnicodeDecodeError:
            try:
                with open(caminho_arquivo, "r", encoding="latin-1") as f:
                    content = f.read()
                    return f"Nome do arquivo: {nome_arquivo}\nConteúdo (parcial): {content[:5000]}"
            except Exception as e:
                print(f"Erro ao ler arquivo de texto '{nome_arquivo}': {e}")
                return f"Arquivo de texto ilegível: {nome_arquivo}"
        except Exception as e:
            print(f"Erro ao ler arquivo de texto '{nome_arquivo}': {e}")
            return f"Arquivo de texto com erro: {nome_arquivo}"

    # Tipos de documentos
    elif extensao == ".docx":
        try:
            document = Document(caminho_arquivo)
            content = "\n".join([para.text for para in document.paragraphs])
            return (
                f"Nome do arquivo: {nome_arquivo}\nConteúdo (parcial): {content[:5000]}"
            )
        except Exception as e:
            print(f"Erro ao ler DOCX '{nome_arquivo}': {e}")
            return f"Documento Word com erro: {nome_arquivo}"
    elif extensao == ".pdf":
        try:
            reader = PdfReader(caminho_arquivo)
            texto = ""
            for page in reader.pages:
                texto += page.extract_text() or ""
            return (
                f"Nome do arquivo: {nome_arquivo}\nConteúdo (parcial): {texto[:5000]}"
            )
        except Exception as e:
            print(f"Erro ao ler PDF '{nome_arquivo}': {e}")
            return f"Documento PDF com erro: {nome_arquivo}"

    # Arquivos de Planilha
    elif extensao == ".xlsx":
        try:
            workbook = openpyxl.load_workbook(caminho_arquivo)
            sheet_names = ", ".join(workbook.sheetnames)
            # Tenta pegar algumas células da primeira aba para dar contexto
            first_sheet = workbook.active
            sample_data = []
            for row_idx in range(
                1, min(first_sheet.max_row + 1, 6)
            ):  # Pega até 5 linhas
                row_values = []
                for col_idx in range(
                    1, min(first_sheet.max_column + 1, 6)
                ):  # Pega até 5 colunas
                    cell_value = first_sheet.cell(row=row_idx, column=col_idx).value
                    if cell_value is not None:
                        row_values.append(str(cell_value))
                if row_values:
                    sample_data.append(", ".join(row_values))

            return (
                f"Nome do arquivo: {nome_arquivo}\n"
                f"Tipo: Planilha Excel\n"
                f"Nomes das abas: {sheet_names}\n"
                f"Dados de exemplo (primeiras linhas/colunas): {'; '.join(sample_data)[:1000]}"
            )  # Limita para IA
        except Exception as e:
            print(f"Erro ao ler XLSX '{nome_arquivo}': {e}")
            return f"Planilha Excel com erro: {nome_arquivo}"
    elif extensao == ".csv":
        try:
            with open(caminho_arquivo, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                sample_lines = []
                for i, row in enumerate(reader):
                    if i >= 10:
                        break
                    sample_lines.append(",".join(row))

            newline_delimiter = "\n"
            return (
                f"Nome do arquivo: {nome_arquivo}\n"
                f"Tipo: Arquivo CSV\n"
                f"Dados de exemplo (primeiras 10 linhas): {newline_delimiter.join(sample_lines)[:5000]}"
            )
        except Exception as e:
            print(f"Erro ao ler CSV '{nome_arquivo}': {e}")
            return f"Arquivo CSV com erro: {nome_arquivo}"

    # Tipos multimídia
    elif extensao in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"):
        return f"Arquivo de imagem: {nome_arquivo}"
    elif extensao in (".mp3", ".wav", ".ogg", ".flac", ".aac"):
        return f"Arquivo de áudio: {nome_arquivo}"
    elif extensao in (".mp4", ".avi", ".mov", ".mkv", ".wmv", ".flv"):
        return f"Arquivo de vídeo: {nome_arquivo}"

    # Outros tipos de arquivo que o Gemini pode categorizar pelo nome ou extensão
    elif extensao in (".zip", ".rar", ".7z", ".tar", ".gz", ".iso"):
        return f"Arquivo compactado ou imagem de disco: {nome_arquivo}"
    elif extensao in (".exe", ".dll", ".msi"):
        return f"Arquivo executável ou de sistema: {nome_arquivo}"

    # Tipo de arquivo não suportado para extração de conteúdo detalhado
    return f"Arquivo não processável por conteúdo: {nome_arquivo}. Extensão: {extensao}"


# --- Função para categorizar arquivo usando Gemini via LangChain ---


def categorizar_arquivo_com_gemini_lc(caminho_arquivo):
    texto_para_analise = extrair_texto_do_arquivo(caminho_arquivo)

    if not texto_para_analise:
        return "Nao_Processado"  # Categoria para arquivos que não puderam ser lidos/analisados

    texto_curto = texto_para_analise[
        :20000
    ]  # Aumentei o limite para mais contexto se necessário

    try:
        categoria = categorization_chain.invoke({"file_content": texto_curto})

        categoria_limpa = "".join(
            c for c in categoria if c.isalnum() or c in [" ", "_", "-"]
        ).strip()
        categoria_limpa = categoria_limpa.replace(" ", "_").replace("__", "_")

        # Garante que a categoria não está vazia após limpeza
        if not categoria_limpa:
            return "Indefinidos"
        return categoria_limpa

    except Exception as e:
        # Se ocorrer um erro de API (limite de taxa, timeout, etc.)
        return "Erro_IA"


# --- Função principal para organizar ---
def organizar_pasta_com_ia_lc(
    caminho_pasta_baguncada, log_func, progress_var, total_files_var
):
    if not os.path.isdir(caminho_pasta_baguncada):
        log_func(
            f"Erro: O caminho '{caminho_pasta_baguncada}' não é um diretório válido.\n"
        )
        return

    log_func(f"Iniciando organização da pasta: {caminho_pasta_baguncada}\n")

    arquivos_a_processar = [
        os.path.join(caminho_pasta_baguncada, item)
        for item in os.listdir(caminho_pasta_baguncada)
        if os.path.isfile(os.path.join(caminho_pasta_baguncada, item))
    ]

    total_files = len(arquivos_a_processar)
    total_files_var.set(total_files)  # Atualiza o total de arquivos na GUI
    processed_count = 0

    for caminho_completo_arquivo in arquivos_a_processar:
        nome_arquivo = os.path.basename(caminho_completo_arquivo)
        log_func(f"\nProcessando: {nome_arquivo}\n")

        categoria_sugerida = categorizar_arquivo_com_gemini_lc(caminho_completo_arquivo)

        if categoria_sugerida:
            pasta_destino = os.path.join(caminho_pasta_baguncada, categoria_sugerida)

            if not os.path.exists(pasta_destino):
                try:
                    os.makedirs(pasta_destino)
                    log_func(f"Pasta criada: {pasta_destino}\n")
                except Exception as e:
                    log_func(
                        f"Erro ao criar pasta {pasta_destino}: {e}. Usando 'Outros_Arquivos'.\n"
                    )
                    pasta_destino = os.path.join(
                        caminho_pasta_baguncada, "Outros_Arquivos"
                    )
                    if not os.path.exists(pasta_destino):
                        os.makedirs(pasta_destino)

            try:
                shutil.move(
                    caminho_completo_arquivo, os.path.join(pasta_destino, nome_arquivo)
                )
                log_func(f"Movido: '{nome_arquivo}' para '{pasta_destino}'\n")
            except shutil.Error as e:
                log_func(
                    f"Erro ao mover '{nome_arquivo}': {e}. Pode ser que já exista um arquivo com esse nome no destino.\n"
                )
            except Exception as e:
                log_func(f"Erro inesperado ao mover '{nome_arquivo}': {e}\n")
        else:
            log_func(
                f"Arquivo '{nome_arquivo}' não foi categorizado. Deixando na pasta original.\n"
            )

        processed_count += 1
        progress_var.set(processed_count)  # Atualiza a barra de progresso

    log_func("\nOrganização concluída!\n")
    # A messagebox deve ser chamada na thread principal da GUI, não na thread de background
    # Então, vamos fazer isso fora desta função, no método _run_organizing


# --- Classe da Interface Gráfica ---
class FileOrganizerApp:
    def __init__(self, master):
        self.master = master
        master.title("Organizador de Arquivos com IA (Gemini)")
        master.geometry("800x600")  # Tamanho inicial da janela

        self.folder_path = tk.StringVar()
        self.progress_var = tk.DoubleVar()
        self.total_files_var = tk.IntVar()

        # --- Frames ---
        self.input_frame = tk.Frame(master, padx=10, pady=10)
        self.input_frame.pack(fill=tk.X)

        self.progress_frame = tk.Frame(master, padx=10, pady=5)
        self.progress_frame.pack(fill=tk.X)

        self.log_frame = tk.Frame(master, padx=10, pady=10)
        self.log_frame.pack(fill=tk.BOTH, expand=True)

        # --- Widgets de Input ---
        tk.Label(self.input_frame, text="Pasta a Organizar:").pack(
            side=tk.LEFT, padx=(0, 5)
        )
        self.folder_entry = tk.Entry(
            self.input_frame, textvariable=self.folder_path, width=60
        )
        self.folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.browse_button = tk.Button(
            self.input_frame, text="Procurar Pasta", command=self.browse_folder
        )
        self.browse_button.pack(side=tk.LEFT, padx=(5, 0))
        self.organize_button = tk.Button(
            self.input_frame,
            text="Organizar!",
            command=self.start_organizing,
            bg="green",
            fg="white",
        )
        self.organize_button.pack(side=tk.LEFT, padx=(10, 0))

        # --- Widgets de Progresso ---
        self.progress_label = tk.Label(
            self.progress_frame, text="Progresso: 0/0 arquivos"
        )
        self.progress_label.pack(side=tk.LEFT, padx=(0, 5))
        self.progressbar = Progressbar(
            self.progress_frame, variable=self.progress_var, maximum=100
        )
        self.progressbar.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Atualiza o label de progresso quando as variáveis mudam
        self.progress_var.trace_add("write", self.update_progress_label)
        self.total_files_var.trace_add("write", self.update_progress_label)

        # --- Log de Saída ---
        tk.Label(self.log_frame, text="Log de Operações:").pack(anchor=tk.W)
        self.output_log = scrolledtext.ScrolledText(
            self.log_frame, wrap=tk.WORD, height=15
        )
        self.output_log.pack(fill=tk.BOTH, expand=True)
        self.output_log.config(
            state=tk.DISABLED
        )  # Começa desabilitado para não permitir digitação

    def browse_folder(self):
        folder_selected = filedialog.askdirectory()
        if folder_selected:
            self.folder_path.set(folder_selected)

    def update_progress_label(self, *args):
        current = int(self.progress_var.get())
        total = self.total_files_var.get()
        self.progress_label.config(text=f"Progresso: {current}/{total} arquivos")
        self.progressbar.config(
            maximum=total if total > 0 else 100
        )  # Evita divisão por zero

    def log_message(self, message):
        self.output_log.config(state=tk.NORMAL)
        self.output_log.insert(tk.END, message)
        self.output_log.see(tk.END)  # Auto-scroll
        self.output_log.config(state=tk.DISABLED)

    def start_organizing(self):
        folder = self.folder_path.get()
        if not folder:
            messagebox.showwarning(
                "Aviso", "Por favor, selecione uma pasta para organizar."
            )
            return

        # Limpa o log e reinicia o progresso
        self.output_log.config(state=tk.NORMAL)
        self.output_log.delete(1.0, tk.END)
        self.output_log.config(state=tk.DISABLED)
        self.progress_var.set(0)
        self.total_files_var.set(0)
        self.organize_button.config(
            state=tk.DISABLED
        )  # Desabilita o botão enquanto processa

        # Inicia a organização em uma thread separada para não travar a GUI
        threading.Thread(target=self._run_organizing, args=(folder,)).start()

    # Dentro da classe FileOrganizerApp:
    def _run_organizing(self, folder):
        try:
            # Passe o método de log da classe, e as variáveis de progresso diretamente
            organizar_pasta_com_ia_lc(
                folder, self.log_message, self.progress_var, self.total_files_var
            )
            messagebox.showinfo(
                "Concluído", "A organização dos arquivos foi finalizada!"
            )  # Mova para aqui
        except Exception as e:
            self.log_message(f"Erro inesperado durante a organização: {e}\n")
            messagebox.showerror("Erro", f"Ocorreu um erro: {e}")
        finally:
            self.organize_button.config(state=tk.NORMAL)  # Reabilita o botão


# --- Exemplo de Uso ---
if __name__ == "__main__":
    # É fundamental definir a variável de ambiente GEMINI_API_KEY antes de executar.
    # Ex: no seu terminal: export GEMINI_API_KEY="SUA_CHAVE_AQUI" (Linux/macOS)
    # Ou: $env:GEMINI_API_KEY="SUA_CHAVE_AQUI" (PowerShell)
    # Ou diretamente no código (apenas para testes rápidos, NÃO RECOMENDADO para produção):
    # os.environ["GEMINI_API_KEY"] = "SUA_CHAVE_API_GEMINI"

    # Crie uma pasta de teste e coloque alguns arquivos variados dentro dela
    # para ver como o script funciona.
    # Ex: 'minha_bagunca_pessoal_gui'

    # pasta_principal_gui_teste = 'minha_bagunca_pessoal_gui'
    # if not os.path.exists(pasta_principal_gui_teste):
    #     os.makedirs(pasta_principal_gui_teste)
    #     with open(os.path.join(pasta_principal_gui_teste, "Relatorio_Vendas_Q1.txt"), "w") as f:
    #         f.write("Este é o relatório de vendas do primeiro trimestre. Detalha receitas e despesas.")
    #     doc = Document()
    #     doc.add_paragraph("Termos e condições do contrato de empréstimo pessoal. Inclui taxas de juros e cronograma de pagamentos.")
    #     doc.save(os.path.join(pasta_principal_gui_teste, "Contrato_Emprestimo.docx"))
    #     with open(os.path.join(pasta_principal_gui_teste, "Foto_Aniversario_Mae.jpg"), "w") as f:
    #         f.write("Conteúdo simulado de uma imagem JPG.")
    #     with open(os.path.join(pasta_principal_gui_teste, "Minha_Musica_Favorita.mp3"), "w") as f:
    #         f.write("Conteúdo simulado de um arquivo MP3.")

    root = tk.Tk()
    app = FileOrganizerApp(root)
    root.mainloop()
